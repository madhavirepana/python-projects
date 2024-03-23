import os
import PyPDF2
import python_docx as docx
from docx import Document
from exception import *
import tkinter as tk
from tkinter import filedialog, messagebox

def split_pdf(pdf_path, output_folder):
    with open(pdf_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page_num in range(len(pdf_reader.pages)):
            pdf_writer = PyPDF2.PdfWriter()
            pdf_writer.add_page(pdf_reader.pages[page_num])
            output_pdf_path = os.path.join(output_folder, f'page_{page_num+1}.pdf')
            with open(output_pdf_path, 'wb') as output_file:
                pdf_writer.write(output_file)

def split_doc(doc_path, output_folder):
    doc = docx.Document(doc_path)
    for i, paragraph in enumerate(doc.paragraphs):
        new_doc = docx.Document()
        new_doc.add_paragraph(paragraph.text)
        output_doc_path = os.path.join(output_folder, f'paragraph_{i+1}.docx')
        new_doc.save(output_doc_path)

def select_pdf():
    pdf_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
    if pdf_path:
        output_folder = filedialog.askdirectory()
        if output_folder:
            split_pdf(pdf_path, output_folder)
            messagebox.showinfo("Success", "PDF successfully split!")

def select_doc():
    doc_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if doc_path:
        output_folder = filedialog.askdirectory()
        if output_folder:
            split_doc(doc_path, output_folder)
            messagebox.showinfo("Success", "DOC successfully split!")

root = tk.Tk()
root.title("File Splitter")

pdf_button = tk.Button(root, text="Select PDF", command=select_pdf)
pdf_button.pack(pady=10)

doc_button = tk.Button(root, text="Select DOC", command=select_doc)
doc_button.pack(pady=10)

root.mainloop()